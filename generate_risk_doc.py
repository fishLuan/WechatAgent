from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Agent \u7ed3\u6784\u5269\u4f59\u98ce\u9669\u5206\u6790\u53ca\u4fee\u590d\u65b9\u6848', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('\u9879\u76ee\uff1aClawBot (WechatAgent)   |   \u65e5\u671f\uff1a2026-07-28   |   \u7248\u672c\uff1av2')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('')

# Section 1
doc.add_heading('\u4e00\u3001\u80cc\u666f\u8bf4\u660e', level=1)
doc.add_paragraph(
    '\u5728\u4e4b\u524d\u7684\u98ce\u9669\u4fee\u590d\u9636\u6bb5\uff0c\u6211\u4eec\u5df2\u4fee\u590d\u4e86 5 \u9879\u9ad8\u98ce\u9669\uff08#5 \u89e6\u9876\u65e0\u53cd\u9988\u3001#11+#6 \u5de5\u5177\u7ed3\u679c\u6ce8\u5165+\u8c0e\u62a5\u3001'
    '#13 \u9650\u6d41\u3001#7 \u6d6a\u8d39 token\uff09\uff0c\u5e76\u901a\u8fc7\u4e86 15/15 \u6d4b\u8bd5\u7528\u4f8b\u3002'
    '\u4ee5\u4e0b\u4e3a\u6587\u6863\u4e2d\u5269\u4f59\u7684 4 \u9879\u98ce\u9669\uff08\u573a\u666f 3/4/6/8\uff09\uff0c\u5c1a\u672a\u5b9e\u65bd\u4fee\u590d\u3002'
)

# Risk data
risks = [
    {
        'id': '\u573a\u666f 3 \u2014 \u5de5\u5177\u6b7b\u9012\u5f52',
        'desc': (
            '\u5de5\u5177 A \u8c03\u7528\u5de5\u5177 B\uff0c\u5de5\u5177 B \u53c8\u8c03\u7528\u5de5\u5177 A\uff0c'
            '\u5f62\u6210\u65e0\u9650\u9012\u5f52\u5faa\u73af\uff0c\u6d88\u8017\u5927\u91cf token \u548c\u65f6\u95f4\uff0c\u6700\u7ec8\u8d85\u65f6\u65e0\u53cd\u9988\u3002'
        ),
        'solution': (
            '\u5728 AgentOrchestrator \u4e2d\u5f15\u5165\u5168\u5c40\u8c03\u7528\u6808 callStack\uff08ThreadLocal\uff09\uff0c'
            '\u6bcf\u6b21\u5de5\u5177\u8c03\u7528\u524d\u68c0\u67e5\u662f\u5426\u5b58\u5728\u5faa\u73af\u4f9d\u8d56\uff08\u540c\u4e00\u5de5\u5177\u91cd\u590d\u5165\u6808\uff09\uff0c'
            '\u8fbe\u5230\u6700\u5927\u6df1\u5ea6\uff085 \u5c42\uff09\u540e\u4e2d\u65ad\u672c\u8f6e\u8c03\u7528\uff0c'
            '\u5f3a\u5236\u6a21\u578b\u57fa\u4e8e\u5df2\u6709\u7ed3\u679c\u56de\u7b54\u3002'
        ),
        'files': [
            'AgentOrchestrator.java \u2014 \u65b0\u589e callStack \u7ba1\u7406\u3001\u6df1\u5ea6\u68c0\u67e5',
            'FunctionCallResult \u2014 \u65b0\u589e INTERRUPTED \u72b6\u6001',
        ],
        'effort': '\u4f4e\uff08\u7ea6 30 \u884c\u4ee3\u7801\uff09',
    },
    {
        'id': '\u573a\u666f 4 \u2014 \u5de5\u5177\u4e92\u76f8\u7529\u9505',
        'desc': (
            '\u5f53\u6240\u6709\u5de5\u5177\u5747\u65e0\u6cd5\u5b8c\u6210\u4efb\u52a1\u65f6\uff08\u5982\u5929\u6c14\u67e5\u4e0d\u5230\u2192\u65b0\u95fb\u4e5f\u67e5\u4e0d\u5230\u2192\u4e92\u76f8\u63a8\u8bf7\uff09\uff0c'
            '\u6a21\u578b\u4f1a\u7f16\u9020\u7406\u7531\u4e0d\u65ad\u91cd\u8bd5\uff0c\u6700\u7ec8\u65e0\u6709\u6548\u56de\u590d\u3002'
        ),
        'solution': (
            '\u5728 AgentOrchestrator \u5916\u5c42\u5faa\u73af\u4e2d\u8bb0\u5f55\u5168\u5c40\u5f02\u5e38\u8ba1\u6570\u3002'
            '\u8fde\u7eed 2 \u6b21\u5de5\u5177\u8c03\u7528\u5931\u8d25\u540e\uff0c'
            '\u7ec8\u6b62\u7f16\u6392\u6d41\u7a0b\u5e76\u56de\u590d\u53cb\u597d\u63d0\u793a\uff08\u201c\u6682\u65e0\u6cd5\u83b7\u53d6\u8be5\u4fe1\u606f\uff0c\u8bf7\u7a0d\u540e\u518d\u8bd5\u201d\uff09\uff0c'
            '\u907f\u514d\u65e0\u9650\u7529\u9505\u3002'
        ),
        'files': [
            'AgentOrchestrator.java \u2014 \u5f02\u5e38\u8ba1\u6570 + \u5168\u5c40\u5157\u5e95\u5206\u652f',
        ],
        'effort': '\u4f4e\uff08\u7ea6 20 \u884c\u4ee3\u7801\uff09',
    },
    {
        'id': '\u573a\u666f 6 \u2014 \u8d85\u957f Content / \u591a\u8f6e\u5806\u79ef',
        'desc': (
            '\u591a\u8f6e\u5bf9\u8bdd\u540e\uff0c\u5386\u53f2\u6d88\u606f\u4e0d\u65ad\u7d2f\u79ef\uff0c'
            '\u5bfc\u81f4 context \u8d85\u957f\uff0c\u6d6a\u8d39 token \u5e76\u53ef\u80fd\u518d\u6b21\u89e6\u53d1\u9650\u989d\u3002'
        ),
        'solution': (
            '\u5728 DeepSeekChatService \u5185\u5c42\u5faa\u73af\u7684\u6bcf\u6b21\u8fed\u4ee3\u524d\uff0c'
            '\u5bf9 messages \u5217\u8868\u505a\u622a\u65ad\uff1a\u4fdd\u7559 System Prompt + \u6700\u8fd1 6 \u8f6e\u7528\u6237/\u52a9\u624b\u5bf9\u8bdd\u3002'
            '\u5f53\u603b token \u4f30\u7b97\u8d85\u8fc7\u9608\u503c\uff08\u5982 8000\uff09\u65f6\uff0c'
            '\u4e22\u5f03\u6700\u65e9\u7684\u975e system \u6d88\u606f\u3002'
        ),
        'files': [
            'DeepSeekChatService.java \u2014 \u6d88\u606f\u622a\u65ad\u903b\u8f91',
            'TokenEstimator.java\uff08\u65b0\u589e\uff09\u2014 \u7b80\u6613 token \u4f30\u7b97\u5de5\u5177\u7c7b',
        ],
        'effort': '\u4e2d\uff08\u7ea6 60 \u884c\u4ee3\u7801\uff09',
    },
    {
        'id': '\u573a\u666f 8 \u2014 \u540c\u4e00\u5de5\u5177\u6b7b\u78d5',
        'desc': (
            '\u5f53\u5de5\u5177\u8fd4\u56de\u65e0\u7ed3\u679c\u65f6\uff0c'
            '\u6a21\u578b\u53ef\u80fd\u7528\u5b8c\u5168\u76f8\u540c\u6216\u6781\u5176\u76f8\u4f3c\u7684\u53c2\u6570\u53cd\u590d\u8c03\u7528\u540c\u4e00\u5de5\u5177\uff0c'
            '\u6d6a\u8d39 token \u4e14\u65e0\u6536\u76ca\u3002'
        ),
        'solution': (
            '\u5728 FunctionToolRegistry \u4e2d\u5f15\u5165\u5de5\u5177\u8c03\u7528\u7f13\u5b58\uff08Caffeine \u6216\u7b80\u5355 ConcurrentHashMap\uff0cTTL 24h\uff09\u3002'
            '\u5bf9\u540c\u4e00\u5de5\u5177\u540d+\u53c2\u6570\u7b7e\u540d\u8fdb\u884c MD5 \u54c8\u5e0c\uff0c'
            '\u7f13\u5b58\u547d\u4e2d\u65f6\u76f4\u63a5\u8fd4\u56de\u4e0a\u6b21\u7ed3\u679c\uff0c\u907f\u514d\u91cd\u590d\u8c03\u7528\u3002'
            '\u7f13\u5b58 key \u542b\u65f6\u95f4\u6233\u7c92\u5ea6\uff08\u5982\u5c0f\u65f6\uff09\uff0c\u907f\u514d\u957f\u671f\u810f\u7f13\u5b58\u3002'
        ),
        'files': [
            'FunctionToolRegistry.java \u2014 \u5de5\u5177\u7ed3\u679c\u7f13\u5b58',
            'application.properties \u2014 \u7f13\u5b58 TTL \u914d\u7f6e\u9879',
        ],
        'effort': '\u4e2d\uff08\u7ea6 50 \u884c\u4ee3\u7801\uff09',
    },
]

for risk in risks:
    doc.add_heading(risk['id'], level=2)

    p = doc.add_paragraph()
    run = p.add_run('\u98ce\u9669\u63cf\u8ff0\uff1a')
    run.bold = True
    p.add_run(risk['desc'])

    p = doc.add_paragraph()
    run = p.add_run('\u4fee\u590d\u65b9\u6848\uff1a')
    run.bold = True
    p.add_run(risk['solution'])

    p = doc.add_paragraph()
    run = p.add_run('\u6d89\u53ca\u6587\u4ef6\uff1a')
    run.bold = True
    for f in risk['files']:
        doc.add_paragraph(f, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('\u9884\u4f30\u5de5\u4f5c\u91cf\uff1a')
    run.bold = True
    run2 = p.add_run(risk['effort'])
    run2.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph('')

# Priority table
doc.add_heading('\u4e09\u3001\u5b9e\u65bd\u4f18\u5148\u7ea7\u5efa\u8bae', level=1)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['\u4f18\u5148\u7ea7', '\u98ce\u9669', '\u5f71\u54cd\u9762', '\u5efa\u8bae']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['P0', '\u573a\u666f 3 \u2014 \u6b7b\u9012\u5f52', '\u7cfb\u7edf\u7a33\u5b9a\u6027 / \u65e0\u9650\u5faa\u73af', '\u672c\u8f6e\u5b9e\u73b0'],
    ['P0', '\u573a\u666f 8 \u2014 \u540c\u4e00\u5de5\u5177\u6b7b\u78d5', 'Token \u6d6a\u8d39 / \u54cd\u5e94\u6162', '\u672c\u8f6e\u5b9e\u73b0'],
    ['P1', '\u573a\u666f 6 \u2014 \u8d85\u957f\u5806\u79ef', 'Token \u6d6a\u8d39', '\u672c\u8f6e\u5b9e\u73b0'],
    ['P1', '\u573a\u666f 4 \u2014 \u4e92\u76f8\u7529\u9505', '\u7528\u6237\u4f53\u9a8c', '\u5982\u679c\u65f6\u95f4\u5141\u8bb8'],
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph('')

# Notes
doc.add_heading('\u56db\u3001\u4f9d\u8d56\u4e0e\u6ce8\u610f\u4e8b\u9879', level=1)
notes = [
    '\u4ee5\u4e0a\u6240\u6709\u4fee\u6539\u5747\u4e3a\u589e\u91cf\u4fee\u6539\uff0c\u4e0d\u5f71\u54cd\u5df2\u6709\u529f\u80fd\u53ca\u5df2\u901a\u8fc7\u7684 15 \u4e2a\u6d4b\u8bd5\u7528\u4f8b\u3002',
    '\u573a\u666f 3 \u7684 callStack \u9700\u4f7f\u7528 ThreadLocal \u786e\u4fdd\u7ebf\u7a0b\u5b89\u5168\u3002',
    '\u573a\u666f 8 \u7684\u7f13\u5b58\u5efa\u8bae\u4f7f\u7528 Caffeine\uff08\u9879\u76ee\u4e2d\u5df2\u6709\u4f9d\u8d56\uff09\uff0c\u907f\u514d\u5f15\u5165\u65b0\u5305\u3002',
    '\u573a\u666f 6 \u7684 token \u4f30\u7b97\u53ef\u91c7\u7528\u7b80\u5316\u65b9\u6848\uff08\u5b57\u7b26\u6570 \u00d7 0.25\uff09\uff0c\u65e0\u9700\u7cbe\u786e tokenizer\u3002',
    '\u5b9e\u65bd\u540e\u5efa\u8bae\u8865\u5145\u5bf9\u5e94\u573a\u666f\u7684\u5355\u5143\u6d4b\u8bd5\uff0c\u8986\u76d6\u5faa\u73af\u68c0\u6d4b/\u7f13\u5b58\u547d\u4e2d/\u622a\u65ad\u8fb9\u754c\u7b49\u3002',
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# Save
path = '/Users/lienqi/Desktop/Agent\u67b6\u6784\u5269\u4f59\u98ce\u9669\u5206\u6790\u4e0e\u4fee\u590d\u65b9\u6848.docx'
doc.save(path)
print('\u6587\u6863\u5df2\u751f\u6210: ' + path)

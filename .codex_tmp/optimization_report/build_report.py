from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Lenovo\Desktop\co-WechatAgent\.codex_tmp\optimization_report\ClawBot基础功能优化与整改方案.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243242"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_RED = "FCE8E6"
PALE_AMBER = "FFF4E5"
PALE_GREEN = "E8F5E9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        p_pr.remove(node)


def add_bottom_border(paragraph, color=BLUE, size="14", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def set_run_font(run, size=None, bold=None, color=None, ascii_font="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_prefix=None, color=None, keep=False):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a, bold=True, color=color or INK)
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b, color=color or INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color or INK)
    if keep:
        set_keep_with_next(p)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run_font(r, size=9, color=INK)
    return table


def add_section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_run_font(r, size=9, bold=True, color=BLUE)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Title", 24, DARK_BLUE, 0, 12),
    ("Subtitle", 13, MUTED, 0, 14),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name not in ("Subtitle",)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    st = styles[list_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.167
    if list_name == "List Bullet 2":
        st.paragraph_format.left_indent = Inches(0.75)
        st.paragraph_format.first_line_indent = Inches(-0.25)
    else:
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)

# Running header
header = section.header
hp = header.paragraphs[0]
hp.paragraph_format.space_after = Pt(0)
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("ClawBot 技术整改方案    ·    内部技术评审")
set_run_font(r, size=8.5, color=MUTED)

# Footer
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("co-WechatAgent  |  ")
set_run_font(r, size=8.5, color=MUTED)
add_field(fp, "PAGE")
for run in fp.runs:
    set_run_font(run, size=8.5, color=MUTED)

# Cover / executive summary
add_section_label(doc, "TECHNICAL REMEDIATION BRIEF")
p = doc.add_paragraph(style="Title")
r = p.add_run("ClawBot 基础功能优化与整改方案")
set_run_font(r, size=24, bold=True, color=DARK_BLUE)
p = doc.add_paragraph(style="Subtitle")
r = p.add_run("代码审查结论、实施路线与验收标准")
set_run_font(r, size=13, color=MUTED)

meta = add_table(
    doc,
    ["项目", "分支", "评审日期", "文档状态"],
    [["co-WechatAgent", "co-luanxu", "2026-07-26", "建议按优先级实施"]],
    [2100, 1900, 2200, 3160],
    header_fill=PALE_BLUE,
)
doc.add_paragraph()

doc.add_heading("执行摘要", level=1)
add_para(
    doc,
    "当前项目的基础调用链已经可以运行，现有 Maven 测试共 16 项并全部通过；但在多用户隔离、网页访问安全、复合任务路由、长消息发送和异常可观测性方面仍存在明显风险。若直接扩大用户量或开放更多工具，问题会从“偶发回答不准确”升级为“跨用户上下文污染、内网探测、消息丢失或故障不可定位”。",
)
add_callout(
    doc,
    "整改结论",
    "优先完成 4 个 P0/P1 项：按用户隔离会话、封堵网页 SSRF、让责任链支持复合任务、统一长文本分片。之后再收紧新闻识别、接入登录超时与统一告警。",
    fill=PALE_AMBER,
    accent="B54708",
)

doc.add_heading("目标结果", level=2)
add_bullet(doc, "不同微信用户之间的历史、摘要、去重集合和清空操作完全隔离。")
add_bullet(doc, "所有外部 URL 访问经过统一安全策略，不可访问本机、内网、链路本地地址或超大响应体。")
add_bullet(doc, "一条消息中的多个独立需求能被拆分、执行并聚合回复，不被第一个处理器提前截断。")
add_bullet(doc, "所有文本回复统一按微信长度限制分片；任何异常都有本地完整日志和钉钉摘要通知。")

page_break(doc)

# Priority overview
doc.add_heading("1. 风险优先级总览", level=1)
add_para(doc, "下表按影响范围、被触发概率和修复收益综合排序。P0 表示应在继续扩展工具前完成；P1 表示进入稳定版本前完成；P2 可与体验优化并行推进。")
priority_rows = [
    ["P0", "会话状态未按用户隔离", "历史串话、清空影响所有人、并发不安全", "引入 userId → ConversationMemory"],
    ["P0", "网页提取存在 SSRF 与内存风险", "可能访问内网/本机，超大响应导致内存压力", "统一 UrlAccessPolicy + 大小/重定向限制"],
    ["P1", "责任链首个命中即停止", "一条消息多个需求只能完成部分", "返回 handled/continue/remainingText"],
    ["P1", "文本回复未统一分片", "聚合答案过长，发送失败或截断", "抽取 WeChatMessageSender"],
    ["P1", "异常被吞且用户看到原始错误", "难排障、可能泄漏内部细节", "集中 ErrorReporter + 安全提示"],
    ["P1", "登录超时配置未生效", "等待循环可能永久阻塞", "实现 deadline、关闭和重试"],
    ["P2", "新闻关键词过宽", "误触发检索、重复上下文、成本增加", "优先 Function Calling 或收窄意图"],
    ["P2", "配置与依赖存在漂移", "假配置、版本不一致、启动负担", "清理默认值、Java 版本和依赖"],
]
add_table(doc, ["级别", "问题", "主要影响", "核心改法"], priority_rows, [720, 2300, 3020, 3320])

doc.add_heading("2. 当前处理流程的关键症结", level=1)
add_callout(
    doc,
    "现状流程",
    "微信消息 → 基础处理器按优先级判断 → 第一个命中的处理器直接结束 → 文本处理器调用规划/模型/工具 → 单次发送回复",
)
add_para(
    doc,
    "这条链路适合“一个输入对应一个动作”的场景，但不适合复合消息。优先级本身没有问题，问题在于处理器只表达“是否处理”，不能表达“已处理一部分、是否继续、还剩哪些文本”。多任务规划即使能拆分，在进入专用处理器之前也可能被责任链提前截断。",
)

doc.add_heading("3. 设计原则", level=1)
add_bullet(doc, "用户隔离优先于模型能力：上下文正确性必须由应用层保证。")
add_bullet(doc, "工具调用与网络访问默认不可信：所有 URL、参数、响应和超时均需边界。")
add_bullet(doc, "编排与执行分离：规划器只拆任务，执行器只运行任务，发送器只负责渠道限制。")
add_bullet(doc, "错误对用户友好、对开发者完整：两类信息分层输出。")

# High priority details
doc.add_heading("4. 高优先级整改项", level=1)

doc.add_heading("4.1 会话记忆按用户隔离", level=2)
add_para(doc, "问题：TextMessageHandler 作为单例持有 longTermSummary、recentMessages 和 processedMsgIds。多个微信用户会共享这些字段；任意用户执行清空操作会影响所有人，而且普通集合在并发回调下不安全。")
add_para(doc, "建议结构：")
add_bullet(doc, "ConversationMemory：保存摘要、最近消息、更新时间和清理锁。")
add_bullet(doc, "ConversationMemoryStore：按稳定的 userId/openId 获取、清空、过期淘汰。")
add_bullet(doc, "MessageDeduplicationService：独立维护消息 ID 的 TTL 有界缓存。")
add_bullet(doc, "持久化文件名包含不可逆用户标识摘要，避免明文暴露并防止用户间覆盖。")
add_callout(doc, "验收点", "A 用户的对话不会出现在 B 用户提示词中；A 清空记忆不影响 B；并发 100 条消息无集合异常；缓存能按 TTL 自动淘汰。", fill=PALE_GREEN, accent="027A48")

doc.add_heading("4.2 网页访问安全与资源边界", level=2)
add_para(doc, "问题：网页提取允许访问 localhost、私网、链路本地地址，并使用不受限的字节数组接收响应。若模型或用户提供恶意 URL，可能形成 SSRF；若响应体过大，会增加内存占用。")
add_para(doc, "统一 UrlAccessPolicy 应至少包含：")
add_bullet(doc, "只允许 http/https；拒绝 localhost、环回、私网、链路本地、保留地址和非常规端口。")
add_bullet(doc, "解析 DNS 后校验所有 IP；每次重定向重新解析并校验，限制最多 3 次。")
add_bullet(doc, "设置连接、读取和总超时；正文上限建议 2 MB，超限立即中止。")
add_bullet(doc, "Content-Type 白名单；日志记录目标域名、耗时和拒绝原因，但不记录敏感查询参数。")
add_callout(doc, "复用要求", "新闻抓取、网页摘要及未来所有联网工具必须共享同一套 URL 访问策略，禁止各自复制校验逻辑。", fill=PALE_RED, accent="B42318")

doc.add_heading("4.3 责任链支持复合任务", level=2)
add_para(doc, "将处理器返回值由简单布尔值改为结构化结果，例如 HandlerResult(handled, continueChain, remainingText, outputs)。专用处理器只消费自己能处理的片段，剩余文本继续进入后续处理器或统一任务执行器。")
add_para(doc, "推荐流程：")
add_number(doc, "先做轻量输入归一化和安全检查。")
add_number(doc, "规划器输出 Task 列表，每项包含类型、原始文本、依赖和顺序。")
add_number(doc, "无依赖任务可并行；有依赖任务串行，并设置单任务超时。")
add_number(doc, "按原问题顺序聚合结果，再交给统一发送器分片。")

# Medium priority
doc.add_heading("5. 稳定性与可观测性整改", level=1)

doc.add_heading("5.1 统一微信消息发送", level=2)
add_para(doc, "目前文档处理存在约 1500 字分片逻辑，而普通文本与多任务聚合结果可能一次发送。应抽取 WeChatMessageSender，统一完成文本清洗、分片、顺序发送、失败重试和发送结果记录。")
add_bullet(doc, "优先按段落边界切分，其次按句号，最后按硬长度切分。")
add_bullet(doc, "每段保留任务编号和上下文；避免把 Markdown 代码块拆成无效片段。")
add_bullet(doc, "单段失败时只重试该段，并使用幂等键防止重复回复。")

doc.add_heading("5.2 统一异常报告", level=2)
add_para(doc, "处理器不应静默吞异常，也不应直接把 e.getMessage() 原样发给用户。建议引入 ErrorReporter：本地日志记录完整堆栈与关联 ID；钉钉发送脱敏摘要；微信只返回可理解的提示和关联 ID。")
error_rows = [
    ["模型/API 超时", "记录模型、耗时、任务类型", "发送告警摘要", "“服务响应超时，请稍后重试”"],
    ["工具参数错误", "记录校验结果", "达到阈值后告警", "指出需补充的字段"],
    ["登录/会话失效", "记录状态转换", "立即告警并附登录入口", "提示管理员处理"],
    ["未知异常", "完整堆栈 + correlationId", "立即告警", "通用失败提示 + 关联 ID"],
]
add_table(doc, ["异常类型", "本地日志", "钉钉通知", "用户提示"], error_rows, [1600, 2600, 2200, 2960])

doc.add_heading("5.3 登录生命周期", level=2)
add_para(doc, "wechat.login.timeout-ms 已有配置入口但等待循环未使用，可能永久阻塞。登录流程应维护明确状态：INIT → QR_READY → WAITING → LOGGED_IN / TIMEOUT / FAILED，并在 deadline 到达后关闭资源、通知钉钉、按退避策略重试。")

doc.add_heading("5.4 新闻识别与多任务上下文", level=2)
add_para(doc, "“今天、昨天、最近、有什么、科技”等词过于宽泛，会把普通问题误判为新闻并提前检索；多任务时还可能把同一新闻上下文复制到每个任务。建议取消宽关键词预取，优先让模型通过 Function Calling 明确选择新闻工具；若保留规则，应要求主题词与时效词同时出现，并只把检索结果绑定到对应任务。")

doc.add_heading("5.5 性能与成本", level=2)
add_para(doc, "多任务规划器不必对每条消息都调用一次模型。可先用轻量门控判断是否包含并列连接词、多个问号或多个已知工具意图；单任务直接走现有链路，复杂输入才进入规划模型。对等价短输入可使用短期规划缓存。")

# Architecture and roadmap
doc.add_heading("6. 目标架构", level=1)
add_callout(
    doc,
    "建议调用链",
    "WeChat Adapter → Input Guard → Task Planner → Task Executor → Tool Registry / LLM → Result Aggregator → WeChatMessageSender",
)
add_para(doc, "横切能力由独立组件提供：ConversationMemoryStore 负责用户上下文；UrlAccessPolicy 约束联网工具；ErrorReporter 统一日志和钉钉通知；ConfigurationProperties 承载模型与第三方 API 配置。")

doc.add_heading("组件职责", level=2)
component_rows = [
    ["WeChat Adapter", "接收/确认消息，提取 userId、msgId、文本和附件", "不包含模型或业务判断"],
    ["Input Guard", "去重、限流、输入长度和基础安全校验", "快速、无模型调用"],
    ["Task Planner", "把复合消息拆为有序任务与依赖", "只规划，不执行"],
    ["Task Executor", "调度工具/LLM，控制并发、超时和取消", "输出结构化结果"],
    ["Tool Registry", "注册工具 schema 与执行器", "工具配置外置"],
    ["Result Aggregator", "按原问题顺序组合成功与失败结果", "保留部分成功"],
    ["Message Sender", "渠道分片、重试、幂等和发送日志", "统一所有回复出口"],
]
add_table(doc, ["组件", "职责", "边界"], component_rows, [1800, 4600, 2960])

doc.add_heading("7. 分阶段实施路线", level=1)
roadmap_rows = [
    ["阶段 1", "数据正确性", "会话隔离、有界去重缓存", "跨用户测试、并发测试通过"],
    ["阶段 2", "网络安全", "UrlAccessPolicy、超时、大小与重定向限制", "SSRF 用例全部拒绝"],
    ["阶段 3", "复合任务", "HandlerResult、规划/执行/聚合拆分", "3 类复合场景全部完成"],
    ["阶段 4", "渠道稳定", "统一分片、重试、关联 ID", "超长回复完整可达"],
    ["阶段 5", "运维体验", "登录超时、钉钉告警、新闻规则收紧", "生命周期与告警测试通过"],
    ["阶段 6", "清理收尾", "配置、依赖和测试覆盖整理", "构建无警告、配置清晰"],
]
add_table(doc, ["阶段", "目标", "交付物", "退出条件"], roadmap_rows, [1200, 1600, 3560, 3000])

add_callout(doc, "并发策略", "只并行执行彼此独立、无共享写入的任务；涉及会话状态、同一文件或具有顺序语义的任务保持串行。每个任务必须有超时，总请求还需设置总体 deadline。", fill=PALE_AMBER, accent="B54708")

# Config & tests
doc.add_heading("8. 配置与依赖清理", level=1)
add_para(doc, "模型、第三方 API、超时和渠道限制应全部通过 @ConfigurationProperties 管理，并在启动时校验必填值。敏感信息只引用环境变量，禁止提交真实密钥。")
config_rows = [
    ["tianapi.api.key", "当前可能使用占位默认值，容易被误判为已配置", "默认空字符串；启用工具时 fail-fast"],
    ["模型配置", "供应商、模型名、温度、超时可能分散在服务类", "集中为 llm.*，按用途区分 planner/chat/tool"],
    ["wechat.login.timeout-ms", "存在但未真正控制等待循环", "绑定到登录 deadline 并覆盖测试"],
    ["钉钉配置", "应避免随分支或本地文件丢失", "只提交变量占位；真实值放环境变量"],
    ["Java 版本", "项目说明为 Java 21，POM 可能仍声明 17", "统一编译、运行与 CI 为 21"],
    ["MySQL 驱动", "当前未见明确使用场景", "无持久化需求则删除"],
    ["Web Starter/8081", "项目无 Controller 时可能是额外启动负担", "确认 SDK 需要后保留，否则设为非 Web"],
    ["NotificationService", "接口上的 @Component 冗余", "仅实现类注册为 Bean"],
]
add_table(doc, ["配置/依赖", "发现", "建议"], config_rows, [2100, 3560, 3700])

doc.add_heading("9. 回归测试矩阵", level=1)
test_rows = [
    ["多用户会话", "A/B 交替聊天、分别清空、并发发送", "上下文完全隔离，无并发异常"],
    ["消息去重", "重复 msgId、过期 msgId、超过容量", "窗口内去重，过期可处理，容量受控"],
    ["复合任务", "天气+汇率、文档+问答、图片+普通文本", "每项均执行，结果按输入顺序聚合"],
    ["URL 安全", "127.0.0.1、私网域名、DNS 重绑定、重定向", "全部拒绝并记录原因"],
    ["响应边界", "超时、2 MB 以上正文、非文本类型", "及时中止，不造成内存峰值"],
    ["长消息", "超过单段限制、多段落、代码块", "完整分片、有序到达、无重复"],
    ["告警", "模型异常、登录超时、未知异常", "本地堆栈完整，钉钉脱敏，用户提示安全"],
    ["配置", "缺少必填密钥、非法超时、生产环境启动", "启动时明确失败或禁用对应工具"],
]
add_table(doc, ["测试域", "关键用例", "期望结果"], test_rows, [1900, 3900, 3560])
add_para(doc, "当前基线：Maven 测试共 16 项并全部通过。该结果说明现有路径未见立即回归，但尚不足以覆盖上述路由、多用户、URL 安全、钉钉告警和登录生命周期场景。")

# Acceptance and decisions
doc.add_heading("10. 验收清单", level=1)
add_section_label(doc, "FUNCTIONAL")
add_bullet(doc, "一条消息包含 2–3 个独立需求时，全部需求都有明确结果；单项失败不吞掉其他成功结果。")
add_bullet(doc, "汇率、天气、文档、图片、新闻等工具均从 Tool Registry 注册并通过统一执行流程调用。")
add_bullet(doc, "所有回复从同一发送组件发出，长文本可稳定分片。")

add_section_label(doc, "SECURITY & RELIABILITY")
add_bullet(doc, "任何联网工具均无法访问本机、私网、链路本地和保留地址。")
add_bullet(doc, "所有外部调用都有连接/读取/总超时和响应大小限制。")
add_bullet(doc, "会话、去重和缓存都具备用户边界、容量边界和时间边界。")
add_bullet(doc, "敏感配置不入库；日志和钉钉消息不包含 API Key、Webhook 签名或完整用户隐私。")

add_section_label(doc, "OPERATIONS")
add_bullet(doc, "登录成功、登录超时、会话失效和未捕获异常均产生符合等级的钉钉通知。")
add_bullet(doc, "每次用户请求具备 correlationId，可贯穿规划、工具调用、聚合、发送和告警。")
add_bullet(doc, "CI 使用 Java 21，clean test 与 clean package 均通过。")

doc.add_heading("11. 关键取舍", level=1)
decision_rows = [
    ["先隔离会话，再扩工具", "避免工具越多，跨用户污染面越大", "需要调整 TextMessageHandler 状态管理"],
    ["统一联网策略", "一次修复覆盖现在和未来工具", "初期需适配现有抓取代码"],
    ["规划器按需调用", "降低延迟和模型成本", "门控规则需持续用真实消息校准"],
    ["部分成功优先", "复合任务中单项失败不影响整体", "聚合结果需明确标注失败项"],
]
add_table(doc, ["决策", "收益", "代价"], decision_rows, [2600, 3900, 2860])

doc.add_heading("12. 完成定义", level=1)
add_callout(
    doc,
    "Definition of Done",
    "上述 P0/P1 项完成；新增回归测试全部通过；配置说明与 application.properties 示例同步；无真实密钥进入 Git；在两个微信用户并发、三任务复合消息、超长回复和恶意 URL 四类验收场景下均表现稳定。",
    fill=PALE_GREEN,
    accent="027A48",
)
add_para(doc, "建议每个阶段独立提交并通过 PR/MR 合并，避免把架构调整、配置迁移和工具新增混在同一提交中，以便审查、回滚和责任归属。", color=MUTED)

# Document metadata
props = doc.core_properties
props.title = "ClawBot 基础功能优化与整改方案"
props.subject = "co-WechatAgent 代码审查与实施路线"
props.author = "Codex"
props.keywords = "ClawBot, WeChat Agent, Spring Boot, 多任务, Function Calling, 安全整改"
props.comments = "基于项目基础功能审查生成"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(str(OUT))
